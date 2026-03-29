import os
import json
import anthropic
from json_repair import repair_json
from docx import Document
from utils import load_yaml, token_usage
from dotenv import load_dotenv
from prompts import WRITER_SYSTEM
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langfuse.decorators import observe, langfuse_context

load_dotenv()

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
config = load_yaml("config.yml")


def load_context(notes_file: str = "notes.txt") -> tuple[str, str]:
    context = {}
    with open(notes_file) as f:
        for line in f:
            if ":" in line:
                key, _, value = line.partition(":")
                context[key.strip()] = value.strip()
    return context["project_folder"], context["topic"]


def load_research_notes(project_folder: str) -> str:
    """Concatenate all task-N.md files into one research brief."""
    notes = []
    i = 1
    while True:
        path = os.path.join(project_folder, f"task-{i}.md")
        if not os.path.exists(path):
            break
        with open(path) as f:
            notes.append(f.read())
        i += 1
    return "\n\n---\n\n".join(notes)


@observe(name="writer", as_type="generation")
def generate_structure(topic: str, research_notes: str) -> dict:
    """Ask the LLM to turn research notes into a document structure."""
    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=8000,
        system=WRITER_SYSTEM,
        messages=[
            {
                "role": "user",
                "content": f"Research topic: {topic}\n\n## Research Notes\n\n{research_notes}",
            }
        ],
    )
    langfuse_context.update_current_observation(
        model=config["model"],
        usage=token_usage(
            config["model"], response.usage.input_tokens, response.usage.output_tokens
        ),
        input=topic,
    )
    text = response.content[0].text.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1]
        text = text.rsplit("```", 1)[0]
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # Response was likely truncated at max_tokens — repair and continue
        return json.loads(repair_json(text))


def render_docx(structure: dict, output_path: str) -> None:
    """Render the document structure to a .docx file."""
    doc = Document()

    # Title
    title = doc.add_heading(structure["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    # Subtitle
    if structure.get("subtitle"):
        sub = doc.add_paragraph(structure["subtitle"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.italic = True
        sub.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()  # spacer

    # Sections
    for section in structure["sections"]:
        doc.add_heading(section["heading"], level=1)
        for block in section["content"]:
            if block["type"] == "paragraph":
                p = doc.add_paragraph(block["text"])
                p.paragraph_format.space_after = Pt(2)
            elif block["type"] == "bullets":
                for item in block["items"]:
                    doc.add_paragraph(item, style="List Bullet")
            # Source line after every block
            source = block.get("source")
            if source:
                is_url = source.startswith("http://") or source.startswith("https://")
                src = doc.add_paragraph(f"Source: {source}")
                src.paragraph_format.space_after = Pt(8)
                run = src.runs[0]
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = (
                    RGBColor(0x80, 0x80, 0x80) if is_url else RGBColor(0x00, 0x00, 0x00)
                )

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(structure["conclusion"])

    doc.save(output_path)
    json_path = output_path.replace(".docx", ".json")
    with open(json_path, "w") as f:
        json.dump(structure, f, indent=2)
    print(f"Document saved: {output_path}")


def write(notes_file: str = "notes.txt") -> None:
    project_folder, topic = load_context(notes_file)
    print(f"Writing document for: '{topic}'")

    research_notes = load_research_notes(project_folder)
    print("Generating document structure...")
    structure = generate_structure(topic, research_notes)

    output_path = os.path.join(project_folder, "report.docx")
    render_docx(structure, output_path)


if __name__ == "__main__":
    write()
